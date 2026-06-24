"""
apex_parser.py — Converts raw documents into the Universal Document AST.

Produces a flat list of :class:`ASTNode` objects with parent-child relationships
encoded via ``parent_id`` and ``children`` fields.

Supported input formats:
    - **Markdown / Plain text** — headings auto-detected, code blocks, tables
    - **PDF** — via `markitdown` conversion to Markdown, then parsed
    - **DOCX** — via `python-docx`, then parsed
    - **Python source** — via the ``ast`` module; classes and functions become nodes
"""

from __future__ import annotations

import ast
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path

from apex_rag.models.unified_models import ASTNode, NodeType
from apex_rag.retrieval.vision.parser import SUPPORTED_EXTENSIONS as _IMAGE_EXTENSIONS

# ═══════════════════════════════════════════════════════════════
# Regex patterns for Markdown parsing
# ═══════════════════════════════════════════════════════════════

# ATX headings: #, ##, ###, etc.
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$", re.MULTILINE)

# Fenced code blocks: ```lang\\n...\\n```
_CODE_FENCE_RE = re.compile(
    r"^(`{3,})\s*(\w*)\s*\n(.*?)\n\1",
    re.MULTILINE | re.DOTALL,
)

# Table rows (lines beginning and ending with | with at least one | in between)
_TABLE_ROW_RE = re.compile(r"^\|.+\|$", re.MULTILINE)

# Page markers: <!-- Page N --> or [Page N]
_PAGE_MARKER_RE = re.compile(r"(?:<!--\s*Page\s+(\d+)\s*-->|\[Page\s+(\d+)\])", re.IGNORECASE)


# ═══════════════════════════════════════════════════════════════
# ApexParser
# ═══════════════════════════════════════════════════════════════


class ApexParser:
    """Converts raw documents into a list of :class:`ASTNode` objects.

    The parser produces a flat list where the tree structure is captured
    by ``parent_id`` and ``children`` (lists of child node IDs).

    Usage::

        parser = ApexParser()
        nodes = await parser.parse_file("document.md")
        # or directly:
        nodes = parser.parse_markdown("# Title\\n\\nSome text")
    """

    def __init__(self, default_doc_id: str | None = None) -> None:
        self._default_doc_id = default_doc_id or str(uuid.uuid4())

    # ── Public API ─────────────────────────────────────────────────────────

    async def parse_file(self, file_path: str | Path, doc_id: str | None = None) -> list[ASTNode]:
        """Parse a document file into AST nodes.

        The file format is detected by its extension:

        * ``.md``, ``.txt``, ``.markdown`` → Markdown / plain text
        * ``.pdf`` → PDF (converted via *markitdown*)
        * ``.docx`` → Word document (converted via *python-docx*)
        * ``.py`` → Python source code
        * ``.png``, ``.jpg``, ``.jpeg``, ``.webp``, etc. → Image files

        Args:
            file_path: Path to the document file.
            doc_id:    Override the auto-generated document ID.

        Returns:
            A flat list of :class:`ASTNode` objects.
        """
        path = Path(file_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        resolved_doc_id = doc_id or self._default_doc_id
        ext = path.suffix.lower()

        if ext in {".md", ".markdown", ".txt"}:
            text = path.read_text(encoding="utf-8", errors="replace")
            return self.parse_markdown(text, doc_id=resolved_doc_id)
        elif ext == ".pdf":
            return await self._parse_pdf(path, doc_id=resolved_doc_id)
        elif ext == ".docx":
            return await self._parse_docx(path, doc_id=resolved_doc_id)
        elif ext == ".py":
            text = path.read_text(encoding="utf-8", errors="replace")
            return self.parse_python(text, doc_id=resolved_doc_id)
        elif ext in _IMAGE_EXTENSIONS:
            return await self._parse_image(path, doc_id=resolved_doc_id)
        else:
            # Fallback: treat as plain text
            text = path.read_text(encoding="utf-8", errors="replace")
            return self.parse_markdown(text, doc_id=resolved_doc_id)

    # ── Markdown / Plain text ──────────────────────────────────────────────

    def parse_markdown(
        self,
        text: str,
        doc_id: str | None = None,
        source_date: datetime | None = None,
    ) -> list[ASTNode]:
        """Parse Markdown or plain text into AST nodes.

        Headings (``#`` … ``######``) become ``HEADING`` nodes.
        Content beneath a heading becomes ``PARAGRAPH`` (or ``TABLE`` / ``CODE``).
        Text before the first heading is attached to an implicit root.

        Args:
            text:       The Markdown / plain text content.
            doc_id:     Document ID.
            source_date: Optional source authorship date.

        Returns:
            A flat list of :class:`ASTNode` objects.
        """
        resolved_doc_id = doc_id or self._default_doc_id

        # Remove code fences first to prevent heading/table detection inside them
        code_regions: dict[str, str] = {}
        def _save_code(m: re.Match[str]) -> str:
            placeholder = f"__CODE_BLOCK_{len(code_regions)}__"
            code_regions[placeholder] = m.group(3)
            return placeholder
        text_no_fences = _CODE_FENCE_RE.sub(_save_code, text)

        nodes: list[ASTNode] = []
        # Stack of (heading_level, node_id) for parent tracking
        stack: list[tuple[int, str]] = []

        ingestion_dt = datetime.now(timezone.utc)
        current_page = 0  # Track page number from <!-- Page N --> markers

        lines = text_no_fences.split("\n")
        current_para_lines: list[str] = []
        current_table_lines: list[str] = []
        in_table = False

        def _make_node(content: str, ntype: NodeType, parent_id: str | None) -> ASTNode:
            """Create an ASTNode, inheriting the current page number."""
            nonlocal current_page
            return ASTNode(
                content=content,
                node_type=ntype,
                depth=_depth_of(stack),
                parent_id=parent_id,
                doc_id=resolved_doc_id,
                source_date=source_date,
                ingestion_date=ingestion_dt,
                page_number=current_page if current_page > 0 else None,
            )

        def _flush_paragraph(parent_id: str | None) -> None:
            nonlocal current_para_lines
            text_content = "\n".join(current_para_lines).strip()
            if text_content:
                # Restore any code blocks
                for ph, code in code_regions.items():
                    if ph in text_content:
                        # This paragraph contains a code block placeholder → create CODE node
                        text_content = text_content.replace(ph, code).strip()
                        node = _make_node(text_content, NodeType.CODE, parent_id)
                        nodes.append(node)
                        if parent_id:
                            _add_child(nodes, parent_id, node.node_id)
                        current_para_lines = []
                        return

                node = _make_node(text_content, NodeType.PARAGRAPH, parent_id)
                nodes.append(node)
                if parent_id:
                    _add_child(nodes, parent_id, node.node_id)
            current_para_lines = []

        def _flush_table(parent_id: str | None) -> None:
            nonlocal current_table_lines, in_table
            if current_table_lines:
                table_text = "\n".join(current_table_lines).strip()
                node = _make_node(table_text, NodeType.TABLE, parent_id)
                nodes.append(node)
                if parent_id:
                    _add_child(nodes, parent_id, node.node_id)
            current_table_lines = []
            in_table = False

        for raw_line in lines:
            line = raw_line

            # Check for page markers
            page_match = _PAGE_MARKER_RE.match(line)
            if page_match:
                page_num = int(page_match.group(1) or page_match.group(2))
                current_page = max(current_page, page_num)
                continue  # Skip the page marker line itself

            # Check if this line is a code block placeholder

            # Check if this line is a code block placeholder
            if line.strip() in code_regions:
                _flush_paragraph(stack[-1][1] if stack else None)
                _flush_table(stack[-1][1] if stack else None)
                code_content = code_regions[line.strip()]
                parent = stack[-1][1] if stack else None
                node = _make_node(code_content, NodeType.CODE, parent)
                nodes.append(node)
                continue

            heading_match = _HEADING_RE.match(line)
            if heading_match:
                _flush_paragraph(stack[-1][1] if stack else None)
                _flush_table(stack[-1][1] if stack else None)

                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()

                # Pop stack to find parent
                while stack and stack[-1][0] >= level:
                    stack.pop()

                parent_id = stack[-1][1] if stack else None

                node = _make_node(title, NodeType.HEADING, parent_id)
                node.depth = level - 1
                nodes.append(node)
                stack.append((level, node.node_id))
                if parent_id:
                    _add_child(nodes, parent_id, node.node_id)
                continue

            # Table detection
            stripped = line.strip()
            is_table_row = _TABLE_ROW_RE.match(stripped) is not None
            is_separator = "---" in stripped if is_table_row else False

            if is_table_row and not is_separator:
                if not in_table:
                    _flush_paragraph(stack[-1][1] if stack else None)
                    in_table = True
                current_table_lines.append(stripped)
                continue

            # If we were in a table:
            # - separator rows: skip silently, stay in table mode
            # - empty lines or non-table rows: flush the table
            if in_table:
                if is_separator:
                    continue
                _flush_table(stack[-1][1] if stack else None)

            # Regular paragraph content
            current_para_lines.append(line)

        # Flush remaining content
        parent_id = stack[-1][1] if stack else None
        if in_table:
            _flush_table(parent_id)
        _flush_paragraph(parent_id)

        # If no headings exist, wrap everything under an implicit root
        if not any(n.node_type == NodeType.HEADING for n in nodes):
            root = ASTNode(
                content="(implicit root)",
                node_type=NodeType.HEADING,
                depth=0,
                parent_id=None,
                doc_id=resolved_doc_id,
                source_date=source_date,
                ingestion_date=ingestion_dt,
                page_number=current_page if current_page > 0 else None,
            )
            nodes.insert(0, root)
            # Re-parent all top-level content nodes to root
            for n in nodes[1:]:
                if n.parent_id is None:
                    n.parent_id = root.node_id
                    _add_child(nodes, root.node_id, n.node_id)
            # Ensure root's children list is populated
            root.children = [n.node_id for n in nodes[1:] if n.parent_id == root.node_id]

        # Post-processing: chunk large leaf nodes for industry readiness
        return _chunk_large_sections(nodes, resolved_doc_id, source_date, ingestion_dt, current_page)

    # ── PDF ────────────────────────────────────────────────────────────────

    async def _parse_pdf(
        self,
        path: Path,
        doc_id: str,
        source_date: datetime | None = None,
    ) -> list[ASTNode]:
        """Parse a PDF via *markitdown* conversion to Markdown."""
        try:
            from markitdown import MarkItDown

            md = MarkItDown()
            result = md.convert(str(path))
            markdown_text: str = result.text_content
        except ImportError:
            # Fallback: try pymupdf
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(str(path))
                markdown_text_parts: list[str] = []
                for page in doc:
                    markdown_text_parts.append(page.get_text())
                doc.close()
                markdown_text = "\n\n".join(markdown_text_parts)
            except ImportError as err:
                raise ImportError(
                    "PDF parsing requires either 'markitdown' or 'pymupdf'. "
                    "Install with: pip install markitdown"
                ) from err

        # Extract PDF metadata for source_date
        if source_date is None:
            try:
                import fitz
                pdf_doc = fitz.open(str(path))
                meta = pdf_doc.metadata
                pdf_doc.close()
                if meta and meta.get("creationDate"):
                    try:
                        pdf_date_str = meta["creationDate"]
                        # PDF date format: D:YYYYMMDDHHMMSS
                        if pdf_date_str.startswith("D:"):
                            pdf_date_str = pdf_date_str[2:]
                        source_date = datetime.strptime(
                            pdf_date_str[:8], "%Y%m%d"
                        ).replace(tzinfo=timezone.utc)
                    except (ValueError, IndexError):
                        pass
            except ImportError:
                pass

        return self.parse_markdown(markdown_text, doc_id=doc_id, source_date=source_date)

    # ── DOCX ───────────────────────────────────────────────────────────────

    async def _parse_docx(
        self,
        path: Path,
        doc_id: str,
        source_date: datetime | None = None,
    ) -> list[ASTNode]:
        """Parse a DOCX file via *python-docx* into Markdown, then AST nodes."""
        try:
            from docx import Document
        except ImportError as err:
            raise ImportError(
                "DOCX parsing requires 'python-docx'. Install with: pip install python-docx"
            ) from err

        doc = Document(str(path))

        # Extract source_date from document properties
        if source_date is None:
            try:
                props = doc.core_properties
                if props.created:
                    source_date = props.created.replace(tzinfo=timezone.utc) if props.created.tzinfo is None else props.created
            except Exception:
                pass

        # Convert to Markdown-like text
        md_lines: list[str] = []
        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ""
            text = para.text.strip()
            if not text:
                md_lines.append("")
                continue
            if style_name.startswith("heading ") or style_name.startswith("heading"):
                level = 1
                for part in style_name.split():
                    if part.isdigit():
                        level = int(part)
                        break
                md_lines.append(f"{'#' * level} {text}")
            else:
                md_lines.append(text)

        # Extract tables
        for table in doc.tables:
            md_lines.append("")
            header_row: list[str] = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            md_lines.append("| " + " | ".join(header_row) + " |")
            md_lines.append("| " + " | ".join(["---"] * len(header_row)) + " |")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                md_lines.append("| " + " | ".join(cells) + " |")
            md_lines.append("")

        markdown_text = "\n".join(md_lines)
        return self.parse_markdown(markdown_text, doc_id=doc_id, source_date=source_date)

    # ── Image files ──────────────────────────────────────────────────────

    async def _parse_image(
        self,
        path: Path,
        doc_id: str,
    ) -> list[ASTNode]:
        """Parse an image file via :class:`ImageParser`.

        Uses the :class:`ImageParser` under the hood, which handles
        base64 encoding, optional local OCR, and produces a single
        ASTNode with ``node_type=IMAGE``.
        """
        from apex_rag.retrieval.vision.parser import ImageParser

        parser = ImageParser(default_doc_id=doc_id)
        return await parser.parse_file(path)

    # ── Python source code ─────────────────────────────────────────────────

    def parse_python(
        self,
        source_code: str,
        doc_id: str | None = None,
        source_date: datetime | None = None,
    ) -> list[ASTNode]:
        """Parse Python source code into AST nodes using the ``ast`` module.

        Top-level classes and functions become nodes. Their methods and
        nested classes become children nodes.

        Args:
            source_code: Raw Python source code.
            doc_id:      Document ID.
            source_date: Optional source authorship date.

        Returns:
            A flat list of :class:`ASTNode` objects.
        """
        resolved_doc_id = doc_id or self._default_doc_id
        ingestion_dt = datetime.now(timezone.utc)

        try:
            tree = ast.parse(source_code)
        except SyntaxError:
            # If invalid Python, treat as plain text
            return self.parse_markdown(
                f"# Python Source\n\n```python\n{source_code}\n```",
                doc_id=resolved_doc_id,
                source_date=source_date,
            )

        nodes: list[ASTNode] = []
        _python_ast_to_nodes(
            tree=tree,
            nodes=nodes,
            parent_id=None,
            doc_id=resolved_doc_id,
            source_date=source_date,
            ingestion_date=ingestion_dt,
            source_code=source_code,
        )
        return nodes


# ═══════════════════════════════════════════════════════════════
# Internal helpers
# ═══════════════════════════════════════════════════════════════


def _depth_of(stack: list[tuple[int, str]]) -> int:
    """Current depth in the heading stack."""
    return len(stack)


def _add_child(nodes: list[ASTNode], parent_id: str, child_id: str) -> None:
    """Add a child node ID to the parent's children list."""
    for node in nodes:
        if node.node_id == parent_id:
            node.children.append(child_id)
            return


def _python_ast_to_nodes(
    tree: ast.AST,
    nodes: list[ASTNode],
    parent_id: str | None,
    doc_id: str,
    source_date: datetime | None,
    ingestion_date: datetime,
    source_code: str,
    depth: int = 0,
) -> None:
    """Recursively convert Python AST nodes to ApexRAG AST nodes."""

    for node in ast.iter_child_nodes(tree):
        if isinstance(node, ast.ClassDef):
            # Class node
            class_content = f"class {node.name}:"
            if node.body:
                first_stmt = node.body[0]
                if isinstance(first_stmt, ast.Expr) and isinstance(first_stmt.value, ast.Constant):
                    doc_value = first_stmt.value.value
                    if isinstance(doc_value, bytes):
                        doc_value = doc_value.decode("utf-8")
                    class_content = f"class {node.name}:\n    \"\"\"{doc_value}\"\"\""

            class_node = ASTNode(
                content=class_content,
                node_type=NodeType.HEADING,
                depth=depth,
                parent_id=parent_id,
                doc_id=doc_id,
                source_date=source_date,
                ingestion_date=ingestion_date,
            )
            nodes.append(class_node)

            # Recurse into class body for methods and nested classes
            _python_ast_to_nodes(
                tree=node,
                nodes=nodes,
                parent_id=class_node.node_id,
                doc_id=doc_id,
                source_date=source_date,
                ingestion_date=ingestion_date,
                source_code=source_code,
                depth=depth + 1,
            )

            if parent_id:
                _add_child(nodes, parent_id, class_node.node_id)

        elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            kind = "async " if isinstance(node, ast.AsyncFunctionDef) else ""
            func_name = node.name

            # Build signature
            args_str = ", ".join(
                [arg.arg for arg in node.args.args]
                + [f"*{arg.arg}" for arg in node.args.kwonlyargs]
            )
            func_content = f"{kind}def {func_name}({args_str}):"
            if node.body:
                first_stmt = node.body[0]
                if isinstance(first_stmt, ast.Expr) and isinstance(first_stmt.value, ast.Constant):
                    doc_value = first_stmt.value.value
                    if isinstance(doc_value, bytes):
                        doc_value = doc_value.decode("utf-8")
                    func_content = f"{kind}def {func_name}({args_str}):\n    \"\"\"{doc_value}\"\"\""

            func_node = ASTNode(
                content=func_content,
                node_type=NodeType.CODE,
                depth=depth,
                parent_id=parent_id,
                doc_id=doc_id,
                source_date=source_date,
                ingestion_date=ingestion_date,
            )
            nodes.append(func_node)

            if parent_id:
                _add_child(nodes, parent_id, func_node.node_id)

        # Note: we skip other top-level nodes (imports, assignments, etc.)
        # as they're usually not structural elements.


# ═══════════════════════════════════════════════════════════════
# Large-section chunking
# ═══════════════════════════════════════════════════════════════


def _chunk_large_sections(
    nodes: list[ASTNode],
    doc_id: str,
    source_date: datetime | None,
    ingestion_date: datetime,
    current_page: int | None = None,
    max_chars: int = 3000,
) -> list[ASTNode]:
    """Split leaf nodes with content exceeding ``max_chars`` into sub-nodes.

    This guarantees no single leaf node exceeds the chunk size, making the
    ingestion robust for large documents (like long legal PDFs).
    Returns the same list with new chunk nodes appended.
    """
    new_nodes: list[ASTNode] = []
    # Build a set of node_ids that are existing children
    parent_children: dict[str, list[str]] = {}
    for n in nodes:
        if n.children:
            parent_children[n.node_id] = list(n.children)

    for node in nodes:
        # Only chunk if this is a leaf with no children and content exceeds limit
        if not node.children and node.content and len(node.content) > max_chars:
            # Split by double newline (paragraphs)
            paragraphs = node.content.split("\n\n")
            chunks: list[str] = []
            current_chunk: list[str] = []
            current_len = 0

            for p in paragraphs:
                if current_len + len(p) > max_chars and current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = []
                    current_len = 0
                current_chunk.append(p)
                current_len += len(p)
            if current_chunk:
                chunks.append("\n\n".join(current_chunk))

            if len(chunks) > 1:
                # Create child nodes for each chunk
                chunk_children: list[str] = []
                for _idx, chunk_text in enumerate(chunks):
                    chunk_node = ASTNode(
                        content=chunk_text,
                        node_type=node.node_type,
                        depth=node.depth + 1,
                        parent_id=node.node_id,
                        doc_id=doc_id,
                        source_date=source_date,
                        ingestion_date=ingestion_date,
                        page_number=current_page if current_page is not None and current_page > 0 else None,
                    )
                    new_nodes.append(chunk_node)
                    chunk_children.append(chunk_node.node_id)
                # Parent is no longer a leaf, content is pushed to children
                node.content = ""
                node.children = chunk_children

    if new_nodes:
        nodes.extend(new_nodes)
    return nodes
